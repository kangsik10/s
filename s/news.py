<<<<<<< HEAD
# from docx import Document

# # 1. 워드 생성하기
# document = Document()

# # 2. 워드 데이터 추가하기
# document.add_heading('기사 제목', level=0)
# document.add_paragraph('기사 링크')
# document.add_paragraph('기사 본문')

# # 3. 워드 저장하기
# document.save("test.docx")
# # [출처] 크롤링 - 파이썬으로 워드문서 다루기 python-docx <삼성전자 뉴스 1개 넣어보기>|작성자 카페인노예



from docx import Document
import requests
from bs4 import BeautifulSoup

# 1. 워드 생성하기
document = Document()

response = requests.get("https://search.naver.com/search.naver?sm=tab_hty.top&where=news&query=%EC%82%BC%EC%84%B1%EC%A0%84%EC%9E%90&oquery=%7Bkeyword%7D&tqi=iaKNxsp0J14ssAv%2B2dhssssss3o-193342")
html = response.text
soup = BeautifulSoup(html,'html.parser')
articles = soup.select("div.info_group")
links = articles[0].select("a.info")
url = links[1].attrs['href']
response = requests.get(url , headers={'User-agent':'Mozila/5.0'})
html = response.text
soup = BeautifulSoup(html , 'html.parser')
title = soup.select_one('#title_area').text
content = soup.select_one("#dic_area")
본문 = content.text


# 2. 워드 데이터 추가하기
document.add_heading(title, level=0)
document.add_paragraph(url)
document.add_paragraph(본문)

# 3. 워드 저장하기
document.save("news.docx")
# [출처] 크롤링 - 파이썬으로 워드문서 다루기 python-docx <삼성전자 뉴스 1개 넣어보기>|작성자 카페인노예
=======
import requests
import csv
from bs4 import BeautifulSoup


for date in range(20230831, 20230800, -1):

    print(date)
    file = open(f"{date}.csv", mode="w", encoding="utf-8", newline="")
    writer = csv.writer(file)

    page = 1

    while True:
        
        print(page)
        
        params = {
            'mode': 'LPOD',
            'oid': '001',
            'date': date,
            'page': page
        }

        headers = {
            'User-Agent':'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/116.0.0.0 Safari/537.36'
        }

        response = requests.get('https://news.naver.com/main/list.naver', headers=headers, params=params)
        bs = BeautifulSoup(response.text, 'html.parser')

        if page != int(bs.select_one('.paging strong').text):
            break

        links = []

        for li in bs.select('ul.type06_headline li'):
            links.append(li.select_one('a').attrs['href'])
        for li in bs.select('ul.type06 li'):
            links.append(li.select_one('a').attrs['href'])

        for link in links:
            response = requests.get(link, headers=headers)
            bs = BeautifulSoup(response.text, 'html.parser')

            title = ''
            content = ''
            category = ''

            if bs.select_one('h2#title_area') != None:
                title = bs.select_one('h2#title_area').text.strip()
                content = bs.select_one('div#contents').text.replace('\n','').strip()
                try:
                    category = bs.select_one('em.media_end_categorize_item').text.strip()
                except:
                    pass

            elif bs.select_one('h4.title') != None:
                title = bs.select_one('h4.title').text.strip()
                content = bs.select_one('div#newsEndContents').text.replace('\n','').strip()
                category = '스포츠'
            elif bs.select_one('h2.end_tit') != None:
                title = bs.select_one('h2.end_tit').text.strip()
                content = bs.select_one('div#articeBody').text.replace('\n','').strip()
                category = '연예'
            else:
                print(link)

            news_id = link.split('/')[-1]
            touch_url = f'https://news.like.naver.com/v1/search/contents?q=JOURNALIST%5B78526(period)%5D%7CNEWS%5Bne_001_{news_id}%5D'
            touch_content = requests.get(touch_url, headers=headers)
            touch_content = touch_content.json()['contents'][-1]['reactions']
            
            all_count = 0

            #스포츠뉴스 처리필요
            for touch in touch_content:
                all_count += touch['count']
                     
            writer.writerow([date, title, category, all_count, content])

        page += 1

print('종료')
>>>>>>> 61bf989f4fa641af31eb148712e479e3f90bb009
